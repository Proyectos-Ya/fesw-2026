"""Adaptadores para cada alternativa de extracción de texto que se evalúa.

Todos exponen la misma interfaz (`Extractor.extraer`) para que el benchmark no
tenga que saber si detrás hay un parser de PDF o un servicio de OCR en la nube.
Agregar un candidato es escribir una subclase y registrarla en `REGISTRO`.

Las importaciones de los extractores del Escalón 1 son perezosas a propósito:
así se puede correr el benchmark con lo que esté instalado, sin exigir todo de
una vez. Un extractor no disponible se reporta como tal en la tabla — que
también es un resultado: "no lo probamos" y "lo probamos y falló" no son lo
mismo.

Nota sobre el orden de la escalera: los del Escalón 1 **no son OCR**. Leen la
capa de texto que el PDF ya trae. Si el documento es digital, ganan siempre —
son exactos, cuestan milisegundos y no tienen dependencias pesadas. Recién
cuando esa capa no existe (un escaneo es una imagen dentro de un PDF) tiene
sentido pagar el costo del Escalón 2, que acá es exclusivamente **OCR en la
nube** — se decidió no evaluar motores locales (Tesseract, PaddleOCR, Apple
Vision), ver `ocr_bench/extractors.py::ExtractorEnLaNube`.
"""

from __future__ import annotations

import os
import time
from abc import ABC, abstractmethod
from dataclasses import dataclass
from pathlib import Path


@dataclass(frozen=True)
class Extraccion:
    """Lo que devuelve cualquier extractor sobre un documento."""

    extractor: str
    paginas: list[str]
    segundos: float
    error: str | None = None

    @property
    def texto(self) -> str:
        return "\n".join(self.paginas)

    @property
    def paginas_vacias(self) -> int:
        """Páginas que salieron sin texto útil.

        Es la forma silenciosa de fallar: el extractor no lanza excepción,
        simplemente devuelve nada para esa página. En un RAG eso es un tramo del
        documento que jamás se va a poder citar, y no se nota mirando el total.
        """
        return sum(1 for p in self.paginas if len(p.strip()) < 10)


class Extractor(ABC):
    """Interfaz común. Una subclase por candidato."""

    nombre: str = "sin-nombre"
    # Cómo clasificarlo en el informe: leer la capa de texto, o servicio de
    # OCR en la nube (que además tiene costo por página).
    familia: str = "texto"
    extensiones: tuple[str, ...] = (".pdf",)

    @abstractmethod
    def disponible(self) -> tuple[bool, str]:
        """¿Se puede correr acá? Devuelve (sí/no, motivo si no)."""

    @abstractmethod
    def _paginas(self, ruta: Path) -> list[str]:
        """Extrae el texto, una entrada por página."""

    def soporta(self, ruta: Path) -> bool:
        return ruta.suffix.lower() in self.extensiones

    def extraer(self, ruta: Path) -> Extraccion:
        """Corre la extracción midiendo el tiempo y sin propagar errores.

        Un extractor que revienta con un documento no debe cortar el benchmark:
        que un motor falle con los escaneados malos es exactamente uno de los
        resultados que se busca, así que se registra y se sigue.
        """
        inicio = time.perf_counter()
        try:
            paginas = self._paginas(ruta)
            error = None
        except Exception as exc:  # noqa: BLE001 - se reporta, no se oculta
            paginas = []
            error = f"{type(exc).__name__}: {exc}"
        return Extraccion(
            extractor=self.nombre,
            paginas=paginas,
            segundos=time.perf_counter() - inicio,
            error=error,
        )


def _falta(modulo: str, paquete: str) -> tuple[bool, str]:
    return False, f"falta el módulo {modulo} (pip install {paquete})"


# --------------------------------------------------------------------------
# Escalón 1 — capa de texto del PDF. Sin OCR, sin modelos, milisegundos.
# --------------------------------------------------------------------------


class PdfplumberExtractor(Extractor):
    """pdfplumber (sobre pdfminer.six). Licencia MIT.

    Elegido para producción entre las alternativas de la capa de texto
    (`PyMuPDF`, `pypdfium2`, `pymupdf4llm` — comparadas y descartadas, ver
    `1.2-ocr-alternativas.md`): mejor fidelidad que `pypdfium2` en el caso
    común (CER 0.024 vs. 0.112 sobre el corpus real) sin la licencia AGPL de
    `PyMuPDF`, que puede obligar a liberar código fuente si el backend se
    distribuye como servicio de red.
    """

    nombre = "pdfplumber"
    familia = "texto"

    def disponible(self) -> tuple[bool, str]:
        try:
            import pdfplumber  # noqa: F401
        except ImportError:
            return _falta("pdfplumber", "pdfplumber")
        return True, ""

    def _paginas(self, ruta: Path) -> list[str]:
        import pdfplumber

        with pdfplumber.open(ruta) as pdf:
            return [pagina.extract_text() or "" for pagina in pdf.pages]


class DocxExtractor(Extractor):
    """python-docx, para los adjuntos .docx del corpus.

    No es OCR ni compite con los demás: un .docx trae el texto en claro y la
    extracción es exacta. Está acá porque el corpus real los incluye
    (`ADJUNTO TÉCNICO.docx`) y el pipeline tiene que cubrirlos.
    """

    nombre = "python-docx"
    familia = "texto"
    extensiones = (".docx",)

    def disponible(self) -> tuple[bool, str]:
        try:
            import docx  # noqa: F401
        except ImportError:
            return _falta("docx", "python-docx")
        return True, ""

    def _paginas(self, ruta: Path) -> list[str]:
        import docx

        documento = docx.Document(str(ruta))
        partes = [p.text for p in documento.paragraphs]
        # Las tablas no están en `paragraphs` y en las bases suelen llevar los
        # plazos y montos. Omitirlas sería perder justo lo que más importa.
        for tabla in documento.tables:
            for fila in tabla.rows:
                partes.append(" | ".join(celda.text for celda in fila.cells))
        # Un .docx no tiene páginas hasta que se renderiza: se devuelve como una
        # sola "página" y el conteo de páginas vacías queda en 0 o 1.
        return ["\n".join(partes)]


# --------------------------------------------------------------------------
# Escalón 2 — OCR en la nube. Se decidió no evaluar OCR local (Tesseract,
# PaddleOCR, Apple Vision): el equipo optó por APIs, empezando por Gemini
# —que el backend ya integra, sin credencial nueva que gestionar— y sumando
# Unstructured y LlamaParse como puntos de comparación.
#
# Nada de esto llama a una red por su cuenta: cada extractor exige sus propias
# variables de entorno (ver cada clase) y sin ellas se reporta "no disponible",
# igual que un motor local sin instalar. La llamada real solo ocurre cuando
# alguien corre el benchmark con las credenciales puestas — nunca al importar
# este módulo ni al listar los extractores.
# --------------------------------------------------------------------------


def _leer_pdf(ruta: Path) -> bytes:
    return ruta.read_bytes()


def _contar_paginas(ruta: Path) -> int:
    """Cuenta páginas con pdfplumber. No es OCR: solo se usa para separar la
    respuesta de una API que devuelve todo el documento en un bloque."""
    import pdfplumber

    with pdfplumber.open(ruta) as documento:
        return len(documento.pages)


class ExtractorEnLaNube(Extractor):
    """Base común: todos cobran por documento/página y necesitan credencial.

    Cada subclase declara qué variables de entorno necesita y las valida en su
    propio `disponible()` — los tres servicios autentican distinto (Gemini:
    API key + nombre de modelo, sin URL propia; Unstructured: API key + URL,
    porque puede ser la nube pública o una instancia propia; LlamaParse: solo
    API key, URL fija) así que forzar una única forma común habría sido más
    confuso que compartir nada.
    """

    familia = "ocr-nube"

    def _falta(self, *variables: str) -> tuple[bool, str]:
        faltantes = [v for v in variables if not os.environ.get(v)]
        if faltantes:
            return False, (
                f"falta configurar {', '.join(faltantes)} "
                "(no se hace ninguna llamada de red sin esto)"
            )
        return True, ""


class GeminiExtractor(ExtractorEnLaNube):
    """Gemini multimodal, leyendo el PDF completo de una vez (no página por página).

    Es el candidato elegido: el backend ya tiene cliente y credencial para
    Gemini (`gemini_deep_analysis_service.py`), mismas variables de entorno
    (`GEMINI_API_KEY`, `GEMINI_MODEL`) — si el backend ya las tiene
    configuradas, este extractor las reutiliza tal cual, sin credencial nueva.

    Gemini 1.5+ acepta un PDF completo como `inline_data` y lo entiende
    nativamente (columnas, tablas, orden de lectura) sin que este arnés tenga
    que rasterizarlo.

    Mejora respecto al patrón del backend (aplicada solo acá, no en
    `gemini_deep_analysis_service.py`)
    ---------------------------------------------------------------------
    La primera versión le pedía a Gemini que insertara un separador de texto
    (`<<<PAGINA>>>`) entre páginas, y partía la respuesta por ese separador.
    Es fràgil: un separador dentro del prompt es una **sugerencia**, no una
    garantía — el modelo puede omitirlo, reformatearlo o perderlo en un
    documento largo, y ahí desaparece la separación por página sin ningún
    aviso de que pasó.

    El backend ya resuelve este tipo de problema de otra forma:
    `generationConfig.responseSchema` fuerza una forma de JSON válida —no es
    una sugerencia, es una restricción estructural sobre cómo decodifica el
    modelo—, y así es como `gemini_deep_analysis_service.py` obtiene
    `compatibility_score`/`recommendation`/`justification` de manera
    confiable. Acá se usa la **misma técnica**, con un schema propio (un
    arreglo de strings, uno por página) en vez de copiar esos campos, que no
    tienen sentido para este caso.

    **Límite no manejado todavía:** `inline_data` tiene un tope de tamaño
    (~20 MB en la API pública). Un PDF más grande necesitaría subirse primero
    por la File API de Gemini, que este extractor no implementa — se
    reportaría como error al toparse con ese límite, no en silencio.

    Riesgo propio: la **alucinación** — puede devolver texto plausible que no
    está en la página. Por eso conviene mirar `token_precision` con más
    atención en este extractor que en los demás; es la métrica que la delata.
    """

    nombre = "gemini"
    familia = "ocr-nube"

    # Códigos de error transitorios de la API (límite de cuota, sobrecarga del
    # modelo) que vale la pena reintentar en vez de reportar como fallo del
    # documento — importan más acá que en las otras tres fuentes porque se va
    # a correr sobre un corpus completo, no un documento suelto.
    _REINTENTOS = 3
    _ESPERA_ENTRE_REINTENTOS = 10
    _CODIGOS_REINTENTABLES = (429, 503)

    def disponible(self) -> tuple[bool, str]:
        return self._falta("GEMINI_API_KEY", "GEMINI_MODEL")

    def _paginas(self, ruta: Path) -> list[str]:
        import base64
        import json
        import time as _time
        import urllib.error
        import urllib.request

        api_key = os.environ["GEMINI_API_KEY"]
        modelo = os.environ["GEMINI_MODEL"]
        total_paginas = _contar_paginas(ruta)

        pdf_b64 = base64.b64encode(_leer_pdf(ruta)).decode("ascii")
        prompt = (
            "Transcribe TEXTUALMENTE y por completo el contenido de este "
            "documento PDF, en español, preservando el orden de lectura real "
            "(si hay columnas, primero la izquierda completa y luego la "
            f"derecha). El documento tiene {total_paginas} páginas. Devuelve "
            "un elemento del arreglo por cada página, en orden, con el texto "
            "completo de esa página. No agregues comentarios, resúmenes ni "
            "texto que no esté en el documento."
        )

        url = (
            "https://generativelanguage.googleapis.com/v1beta/models/"
            f"{modelo}:generateContent"
        )
        payload = {
            "contents": [
                {
                    "parts": [
                        {"inline_data": {"mime_type": "application/pdf", "data": pdf_b64}},
                        {"text": prompt},
                    ]
                }
            ],
            # Salida forzada a JSON con esta forma exacta: no depende de que
            # el modelo "decida" seguir el formato pedido en el prompt.
            "generationConfig": {
                "responseMimeType": "application/json",
                "responseSchema": {
                    "type": "ARRAY",
                    "items": {"type": "STRING"},
                },
            },
        }
        cuerpo_peticion = json.dumps(payload).encode("utf-8")

        ultimo_error: Exception | None = None
        for intento in range(self._REINTENTOS):
            peticion = urllib.request.Request(
                url,
                data=cuerpo_peticion,
                method="POST",
                headers={
                    "Content-Type": "application/json",
                    # Google dejó de aceptar el parámetro `?key=` para las keys
                    # nuevas de AI Studio ("auth keys", ligadas a una cuenta de
                    # servicio) — hay que mandarla en este header. Las keys
                    # "Standard" antiguas también lo aceptan, así que este
                    # cambio no rompe compatibilidad hacia atrás.
                    "x-goog-api-key": api_key,
                },
            )
            try:
                with urllib.request.urlopen(peticion, timeout=180) as respuesta:
                    cuerpo = json.loads(respuesta.read().decode("utf-8"))
                break
            except urllib.error.HTTPError as exc:
                ultimo_error = exc
                if exc.code not in self._CODIGOS_REINTENTABLES or intento == self._REINTENTOS - 1:
                    raise
                _time.sleep(self._ESPERA_ENTRE_REINTENTOS * (intento + 1))
        else:  # pragma: no cover - inalcanzable, el raise de arriba corta antes
            raise ultimo_error  # type: ignore[misc]

        texto_json = cuerpo["candidates"][0]["content"]["parts"][0]["text"]
        paginas = json.loads(texto_json)
        if not isinstance(paginas, list):
            # La API garantiza la forma vía responseSchema; esto es una
            # defensa ante un cambio de comportamiento futuro, no el camino
            # esperado.
            raise ValueError(
                f"Gemini no devolvió un arreglo pese al responseSchema: {texto_json[:200]!r}"
            )
        return [str(p) for p in paginas]


class DocumentAIExtractor(ExtractorEnLaNube):
    """Google Document AI. **No es generativo**: reconoce texto token por
    token, con posición y confianza por palabra, en vez de "redactar" una
    transcripción como hace un modelo como Gemini. Es la comparación que
    responde la pregunta real de esta ronda: ¿vale la pena la fidelidad de un
    OCR clásico frente a la conveniencia de un modelo generativo, en
    documentos con fechas y montos donde una alucinación silenciosa sale cara?

    Autenticación distinta a las otras tres, y **es la parte más frágil de
    este extractor**: Document AI usa OAuth2 de Google Cloud, no una API key
    plana. Para no traer el SDK de `google-cloud-documentai` (dependencia
    pesada, con su propio manejo de credenciales) a un PoC, este extractor
    espera un **token de acceso ya generado**:

        gcloud auth print-access-token

    Ese token **expira en aproximadamente una hora**. Para una corrida larga
    del benchmark hay que regenerarlo y volver a exportarlo — a diferencia de
    Gemini/Unstructured/LlamaParse, cuya API key no vence. Si en algún momento
    esto se vuelve incómodo, migrar a una cuenta de servicio con
    `google-auth` es el camino correcto, pero no se justifica para probar.

    Requiere haber creado antes un procesador de tipo "OCR" (o "Document OCR")
    en Google Cloud Console → Document AI, en algún `location` (`us` o `eu`).
    """

    nombre = "document-ai"
    familia = "ocr-nube"

    def disponible(self) -> tuple[bool, str]:
        return self._falta(
            "DOCUMENT_AI_PROJECT_ID",
            "DOCUMENT_AI_LOCATION",
            "DOCUMENT_AI_PROCESSOR_ID",
            "DOCUMENT_AI_ACCESS_TOKEN",
        )

    def _paginas(self, ruta: Path) -> list[str]:
        import base64
        import json
        import urllib.request

        proyecto = os.environ["DOCUMENT_AI_PROJECT_ID"]
        ubicacion = os.environ["DOCUMENT_AI_LOCATION"]
        procesador = os.environ["DOCUMENT_AI_PROCESSOR_ID"]
        token = os.environ["DOCUMENT_AI_ACCESS_TOKEN"]

        url = (
            f"https://{ubicacion}-documentai.googleapis.com/v1/projects/"
            f"{proyecto}/locations/{ubicacion}/processors/{procesador}:process"
        )
        payload = {
            "rawDocument": {
                "content": base64.b64encode(_leer_pdf(ruta)).decode("ascii"),
                "mimeType": "application/pdf",
            }
        }
        peticion = urllib.request.Request(
            url,
            data=json.dumps(payload).encode("utf-8"),
            method="POST",
            headers={
                "Authorization": f"Bearer {token}",
                "Content-Type": "application/json; charset=utf-8",
            },
        )
        with urllib.request.urlopen(peticion, timeout=180) as respuesta:
            cuerpo = json.loads(respuesta.read().decode("utf-8"))

        documento = cuerpo["document"]
        texto_completo = documento.get("text", "")
        paginas_crudas = documento.get("pages")
        if not paginas_crudas:
            return [texto_completo]

        # El texto viene como un solo bloque; cada página solo trae los
        # rangos (`textSegments`) que le corresponden dentro de ese bloque.
        # Reconstruir la página es recortar esos rangos y unirlos.
        paginas: list[str] = []
        for pagina in paginas_crudas:
            segmentos = (
                (pagina.get("layout") or {}).get("textAnchor") or {}
            ).get("textSegments") or []
            trozos = [
                texto_completo[int(seg.get("startIndex", 0)) : int(seg["endIndex"])]
                for seg in segmentos
                if "endIndex" in seg
            ]
            paginas.append("".join(trozos) if trozos else "")
        return paginas


class UnstructuredExtractor(ExtractorEnLaNube):
    """Unstructured, vía la **Workflow API** (jobs asíncronos), no la API
    clásica de un solo POST.

    La cuenta usada para este PoC viene aprovisionada para la Workflow API
    (`platform-api.transform.unstructured.io`) — la API clásica
    (`api.unstructuredapp.io/general/v0/general`, un solo POST con el archivo
    y la respuesta inmediata) devuelve 404 ahí, porque es un producto
    distinto. La Workflow API no tiene un contrato REST público documentado
    (Unstructured remite a su SDK y a Swagger UI con la propia API key), así
    que se usa el SDK oficial (`unstructured-client`) en vez de armar las
    llamadas a mano — a diferencia del resto de este archivo, que evita
    dependencias nuevas cuando el contrato es simple y público.

    Tres pasos, como LlamaParse: crear el job, hacer polling hasta que
    termine, descargar el resultado. La diferencia es que acá el job se
    define ad-hoc en la propia petición (`job_nodes`), sin depender de un
    workflow guardado de antemano en el dashboard — no hace falta configurar
    nada ahí para correr esto.

    Nodo de partición: `hi_res` (subtype `unstructured_api` en el payload del
    job) — el motor propio de Unstructured (detección de layout + OCR
    clásico), no el nodo `vlm` (que reenvía la imagen a Claude por defecto y
    puede consumir cuota/crédito aparte). Mismo criterio que ya se usaba en
    la API clásica antes de este cambio.
    """

    nombre = "unstructured"
    familia = "ocr-nube"

    _INTENTOS_MAXIMOS = 60
    _ESPERA_ENTRE_INTENTOS = 5

    def disponible(self) -> tuple[bool, str]:
        return self._falta("UNSTRUCTURED_API_URL", "UNSTRUCTURED_API_KEY")

    def _paginas(self, ruta: Path) -> list[str]:
        import json
        import mimetypes
        import time as _time

        from unstructured_client import UnstructuredClient
        from unstructured_client.models.operations import (
            CreateJobRequest,
            DownloadJobOutputRequest,
            GetJobRequest,
        )
        from unstructured_client.models.shared import BodyCreateJob, InputFiles, JobStatus

        cliente = UnstructuredClient(
            api_key_auth=os.environ["UNSTRUCTURED_API_KEY"],
            server_url=os.environ["UNSTRUCTURED_API_URL"],
        )

        content_type, _ = mimetypes.guess_type(str(ruta))
        with open(ruta, "rb") as archivo:
            respuesta_job = cliente.jobs.create_job(
                request=CreateJobRequest(
                    body_create_job=BodyCreateJob(
                        request_data=json.dumps(
                            {
                                "job_nodes": [
                                    {
                                        "name": "Partitioner",
                                        "type": "partition",
                                        "subtype": "unstructured_api",
                                        "settings": {
                                            "strategy": "hi_res",
                                            "ocr_languages": ["spa"],
                                        },
                                    }
                                ]
                            }
                        ),
                        input_files=[
                            InputFiles(
                                content=archivo,
                                file_name=ruta.name,
                                content_type=content_type or "application/octet-stream",
                            )
                        ],
                    )
                )
            )

        job_id = respuesta_job.job_information.id

        info = None
        for _ in range(self._INTENTOS_MAXIMOS):
            info = cliente.jobs.get_job(request=GetJobRequest(job_id=job_id)).job_information
            if info.status in (JobStatus.COMPLETED, JobStatus.FAILED, JobStatus.STOPPED):
                break
            _time.sleep(self._ESPERA_ENTRE_INTENTOS)
        else:
            raise TimeoutError(
                f"Unstructured: el job {job_id} no terminó tras "
                f"{self._INTENTOS_MAXIMOS * self._ESPERA_ENTRE_INTENTOS}s"
            )

        if info is None or info.status != JobStatus.COMPLETED:
            estado = info.status.value if info else "desconocido"
            raise RuntimeError(f"Unstructured: job {job_id} terminó en estado {estado}")

        por_pagina: dict[int, list[str]] = {}
        for archivo_salida in info.output_node_files or []:
            resultado = cliente.jobs.download_job_output(
                request=DownloadJobOutputRequest(
                    job_id=job_id,
                    file_id=archivo_salida.file_id,
                    node_id=archivo_salida.node_id,
                )
            )
            elementos = resultado.any
            if not isinstance(elementos, list):
                continue
            for elemento in elementos:
                numero = int((elemento.get("metadata") or {}).get("page_number") or 1)
                texto = str(elemento.get("text", ""))
                por_pagina.setdefault(numero, []).append(texto)

        if not por_pagina:
            return []
        return ["\n".join(por_pagina[n]) for n in sorted(por_pagina)]


class LlamaParseExtractor(ExtractorEnLaNube):
    """LlamaParse (LlamaCloud). Pensado explícitamente para RAG: la salida es
    Markdown con la estructura del documento conservada.

    Es una API **asíncrona por trabajo** (`job`), a diferencia de Gemini y
    Unstructured que responden en la misma llamada: se sube el documento, se
    consulta el estado hasta que termina, y recién ahí se pide el resultado.
    El costo de eso es la latencia — puede tardar bastante más por documento
    que los otros dos — y es justo lo que hay que medir con `s/pág` en la
    tabla del benchmark.

    **Sin verificar en vivo todavía**, igual que Unstructured: los nombres de
    endpoint y de campos son los documentados públicamente por LlamaCloud, no
    confirmados contra una respuesta real.
    """

    nombre = "llamaparse"
    familia = "ocr-nube"

    _URL_BASE = "https://api.cloud.llamaindex.ai/api/parsing"
    _INTENTOS_MAXIMOS = 60
    _ESPERA_ENTRE_INTENTOS = 5

    def disponible(self) -> tuple[bool, str]:
        return self._falta("LLAMA_CLOUD_API_KEY")

    def _paginas(self, ruta: Path) -> list[str]:
        import json
        import time as _time
        import urllib.error
        import urllib.request
        import uuid

        token = os.environ["LLAMA_CLOUD_API_KEY"]
        cabeceras = {"Authorization": f"Bearer {token}"}

        limite = uuid.uuid4().hex
        cuerpo = _codificar_multipart(limite, {"language": "es"}, "file", ruta)
        subida = urllib.request.Request(
            f"{self._URL_BASE}/upload",
            data=cuerpo,
            method="POST",
            headers={
                **cabeceras,
                "Content-Type": f"multipart/form-data; boundary={limite}",
                "Accept": "application/json",
            },
        )
        with urllib.request.urlopen(subida, timeout=60) as respuesta:
            trabajo = json.loads(respuesta.read().decode("utf-8"))
        id_trabajo = trabajo["id"]

        # Espera activa con límite: la API es asíncrona y no hay forma de que
        # avise cuando termina, solo se puede preguntar. Un límite de intentos
        # evita que un trabajo colgado deje el benchmark esperando para siempre.
        for _ in range(self._INTENTOS_MAXIMOS):
            estado_peticion = urllib.request.Request(
                f"{self._URL_BASE}/job/{id_trabajo}", headers=cabeceras
            )
            with urllib.request.urlopen(estado_peticion, timeout=30) as respuesta:
                estado = json.loads(respuesta.read().decode("utf-8"))
            if estado.get("status") == "SUCCESS":
                break
            if estado.get("status") == "ERROR":
                raise RuntimeError(f"LlamaParse reportó error en el trabajo: {estado}")
            _time.sleep(self._ESPERA_ENTRE_INTENTOS)
        else:
            raise TimeoutError(
                f"LlamaParse no terminó tras {self._INTENTOS_MAXIMOS * self._ESPERA_ENTRE_INTENTOS}s"
            )

        try:
            resultado_peticion = urllib.request.Request(
                f"{self._URL_BASE}/job/{id_trabajo}/result/json", headers=cabeceras
            )
            with urllib.request.urlopen(resultado_peticion, timeout=60) as respuesta:
                resultado = json.loads(respuesta.read().decode("utf-8"))
            paginas = resultado.get("pages")
            if paginas:
                return [str(p.get("md") or p.get("text") or "") for p in paginas]
        except (urllib.error.HTTPError, KeyError, ValueError):
            pass  # cae al resultado en un solo bloque, más abajo

        # Respaldo: el resultado en Markdown de una vez, sin separación por
        # página. Sirve igual para CER/recall del documento completo; solo se
        # pierde el detalle de páginas vacías.
        respaldo_peticion = urllib.request.Request(
            f"{self._URL_BASE}/job/{id_trabajo}/result/markdown", headers=cabeceras
        )
        with urllib.request.urlopen(respaldo_peticion, timeout=60) as respuesta:
            cuerpo_md = json.loads(respuesta.read().decode("utf-8"))
        return [str(cuerpo_md.get("markdown", ""))]


def _codificar_multipart(
    limite: str, campos: dict[str, str], nombre_archivo: str, ruta: Path
) -> bytes:
    """Arma un cuerpo `multipart/form-data` a mano, sin librerías de terceros.

    `urllib` (stdlib) no trae un cliente HTTP con soporte de multipart —
    `requests` sí, pero es una dependencia nueva solo para esto. Es un formato
    simple de armar directamente: un bloque de texto por campo, uno con el
    archivo, todos separados por el mismo `limite`.
    """
    partes = []
    for clave, valor in campos.items():
        partes.append(
            f"--{limite}\r\n"
            f'Content-Disposition: form-data; name="{clave}"\r\n\r\n'
            f"{valor}\r\n".encode("utf-8")
        )
    partes.append(
        (
            f"--{limite}\r\n"
            f'Content-Disposition: form-data; name="{nombre_archivo}"; filename="{ruta.name}"\r\n'
            "Content-Type: application/pdf\r\n\r\n"
        ).encode("utf-8")
        + ruta.read_bytes()
        + b"\r\n"
    )
    partes.append(f"--{limite}--\r\n".encode("utf-8"))
    return b"".join(partes)


REGISTRO: dict[str, Extractor] = {
    extractor.nombre: extractor
    for extractor in (
        PdfplumberExtractor(),
        DocxExtractor(),
        GeminiExtractor(),
        DocumentAIExtractor(),
        UnstructuredExtractor(),
        LlamaParseExtractor(),
    )
}


def disponibles() -> dict[str, str]:
    """Diagnóstico: qué extractor está listo y por qué no lo está el resto."""
    return {
        nombre: "" if ok else motivo
        for nombre, extractor in REGISTRO.items()
        for ok, motivo in [extractor.disponible()]
    }
